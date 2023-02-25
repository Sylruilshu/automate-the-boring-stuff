import docx


doc = docx.Document("custom_invitations.docx")
text_file = open("guests.txt", "r")
guests = text_file.readlines()
text_file.close()

for guest in guests:
    doc.add_paragraph("It would be a pleasure to have the company of")
    doc.add_paragraph(guest.strip("\n"))
    doc.add_paragraph("at 11010 Memory Lane on the evening of")
    doc.add_paragraph("April 1st")
    doc.add_paragraph("at 7 o'clock")

    doc.add_page_break()
    # doc.paragraphs[5].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)

doc.save("custom_invitations.docx")
